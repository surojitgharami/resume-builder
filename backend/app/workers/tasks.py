# app/workers/tasks.py
import logging
from typing import Dict, Any, Optional
from datetime import datetime, timedelta
import asyncio
from functools import wraps

from app.workers.celery_app import celery_app
from app.models.resume import ResumeStatus, ResumeCreate, ResumeFormat
from app.core.config import settings

logger = logging.getLogger(__name__)


def async_task(func):
    """Decorator to run async functions in Celery tasks."""
    @wraps(func)
    def wrapper(*args, **kwargs):
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # If already in an event loop, create a new one
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
        return loop.run_until_complete(func(*args, **kwargs))
    return wrapper


def get_task_db():
    """
    Get database connection for Celery tasks.
    
    Uses the worker process database connection pool.
    This is more efficient than creating a new client per task.
    """
    from app.workers.celery_app import get_worker_db
    return get_worker_db()


@celery_app.task(name="generate_resume_async", bind=True, max_retries=3)
@async_task
async def generate_resume_async(self, user_id: str, resume_id: str, request_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Background task for generating resumes asynchronously.
    
    Args:
        user_id: User ID
        resume_id: Resume ID (pre-created)
        request_data: Resume generation request data
        
    Returns:
        Task result dictionary
    """
    db = get_task_db()
    
    try:
        logger.info(f"Starting async resume generation for resume_id={resume_id}, user_id={user_id}")
        
        # Get worker services (thread-safe, per-process instances)
        from app.workers.worker_services import (
            get_worker_llm_service,
            get_worker_embeddings_service,
            get_worker_pdf_service,
            get_worker_storage_service
        )
        from app.services.resume_generator import ResumeGeneratorService
        from app.models.user import User
        from io import BytesIO
        
        llm_service = get_worker_llm_service()
        embeddings_service = get_worker_embeddings_service()
        pdf_generator_service = get_worker_pdf_service()
        storage_service = get_worker_storage_service()
        
        # Mark resume as processing
        await db["resumes"].update_one(
            {"resume_id": resume_id},
            {
                "$set": {
                    "status": ResumeStatus.PROCESSING.value,
                    "metadata.task_id": self.request.id,
                    "metadata.started_at": datetime.utcnow()
                }
            }
        )
        
        # Fetch user from database
        user_data = await db["users"].find_one({"_id": user_id})
        if not user_data:
            raise Exception(f"User not found: {user_id}")
        
        user = User(**user_data)
        
        # Parse request data
        resume_request = ResumeCreate(**request_data)
        
        # Initialize services
        generator_service = ResumeGeneratorService(
            db=db,
            llm_service=llm_service,
            embeddings_service=embeddings_service,
            pdf_service=pdf_generator_service,
            storage_service=storage_service
        )
        
        # Fetch the resume record
        resume_data = await db["resumes"].find_one({"resume_id": resume_id})
        if not resume_data:
            raise Exception(f"Resume record not found: {resume_id}")
        
        from app.models.resume import Resume
        resume = Resume(**resume_data)
        
        # Get RAG context if enabled
        context = None
        if resume_request.use_rag:
            try:
                from app.services.rag import RAGService
                rag_service = RAGService(db, embeddings_service)
                results = await rag_service.search_similar(
                    user_id=user_id,
                    query=resume_request.job_description,
                    top_k=5
                )
                if results:
                    context = "\n\n".join([r.get("content", "") for r in results])
                    logger.info(f"Retrieved {len(results)} RAG documents for context")
            except Exception as e:
                logger.warning(f"RAG context retrieval failed: {e}")
        
        # Generate resume sections
        sections = await generator_service._generate_sections(
            user=user,
            job_description=resume_request.job_description,
            template_preferences=resume_request.template_preferences,
            context=context,
            custom_instructions=resume_request.custom_instructions
        )
        
        resume.sections = sections
        resume.status = ResumeStatus.COMPLETED
        resume.completed_at = datetime.utcnow()
        
        # Generate PDF if requested
        if resume_request.format == ResumeFormat.PDF:
            try:
                pdf_bytes = await pdf_generator_service.generate_pdf(resume)
                
                # Upload to S3
                s3_key = f"resumes/{user_id}/{resume_id}.pdf"
                await storage_service.upload_file(
                    BytesIO(pdf_bytes),
                    s3_key,
                    content_type="application/pdf",
                    metadata={
                        "user_id": user_id,
                        "resume_id": resume_id,
                        "generated_at": datetime.utcnow().isoformat()
                    }
                )
                
                # Generate presigned URL
                download_url = await storage_service.generate_presigned_url(
                    s3_key,
                    expiration=7200  # 2 hours
                )
                
                resume.s3_key = s3_key
                resume.download_url = download_url
                logger.info(f"PDF generated and uploaded to S3: {s3_key}")
            except Exception as e:
                logger.warning(f"PDF generation failed, resume saved as JSON: {e}")
        
        # Update resume in database
        await db["resumes"].update_one(
            {"resume_id": resume_id},
            {"$set": resume.dict()}
        )
        
        logger.info(f"Successfully completed async resume generation for resume_id={resume_id}")
        
        return {
            "status": "success",
            "resume_id": resume_id,
            "user_id": user_id,
            "sections_count": len(sections),
            "completed_at": datetime.utcnow().isoformat(),
            "format": resume_request.format.value,
            "download_url": resume.download_url
        }
        
    except Exception as e:
        logger.error(f"Resume generation task failed for resume_id={resume_id}: {e}", exc_info=True)
        
        # Update resume status to failed
        try:
            await db["resumes"].update_one(
                {"resume_id": resume_id},
                {
                    "$set": {
                        "status": ResumeStatus.FAILED.value,
                        "error_message": str(e),
                        "metadata.failed_at": datetime.utcnow(),
                        "metadata.task_id": self.request.id
                    }
                }
            )
        except Exception as update_error:
            logger.error(f"Failed to update resume status: {update_error}")
        
        # Retry logic
        if self.request.retries < self.max_retries:
            logger.info(f"Retrying task (attempt {self.request.retries + 1}/{self.max_retries})")
            raise self.retry(exc=e, countdown=60 * (self.request.retries + 1))
        
        return {
            "status": "failed",
            "resume_id": resume_id,
            "user_id": user_id,
            "error": str(e),
            "failed_at": datetime.utcnow().isoformat()
        }


@celery_app.task(name="process_uploaded_resume", bind=True, max_retries=3)
@async_task
async def process_uploaded_resume(self, file_id: str, user_id: str) -> Dict[str, Any]:
    """
    Background task for processing uploaded resume files (OCR and data extraction).
    
    Args:
        file_id: File upload ID
        user_id: User ID
        
    Returns:
        Task result dictionary
    """
    db = get_task_db()
    
    try:
        logger.info(f"Processing uploaded resume file_id={file_id}, user_id={user_id}")
        
        # Get worker services (thread-safe, per-process instances)
        from app.workers.worker_services import (
            get_worker_ocr_service,
            get_worker_llm_service,
            get_worker_storage_service,
            get_worker_embeddings_service
        )
        from app.services.rag import RAGService
        
        ocr_service = get_worker_ocr_service()
        llm_service = get_worker_llm_service()
        storage_service = get_worker_storage_service()
        embeddings_service = get_worker_embeddings_service()
        
        # Fetch file upload record
        file_data = await db["uploads"].find_one({"file_id": file_id, "user_id": user_id})
        if not file_data:
            raise Exception(f"File upload not found: {file_id}")
        
        # Mark as processing
        await db["uploads"].update_one(
            {"file_id": file_id},
            {
                "$set": {
                    "metadata.processing": True,
                    "metadata.task_id": self.request.id,
                    "metadata.processing_started_at": datetime.utcnow()
                }
            }
        )
        
        s3_key = file_data["s3_key"]
        filename = file_data["filename"]
        file_ext = filename.split('.')[-1].lower()
        
        logger.info(f"Downloading file from S3: {s3_key}")
        
        # 1. Download file from S3
        file_bytes = await storage_service.download_file(s3_key)
        
        # 2. Run OCR if needed
        ocr_text = None
        if file_ext in ['pdf']:
            logger.info(f"Running OCR on PDF: {filename}")
            try:
                ocr_text = await ocr_service.extract_text_from_pdf(file_bytes)
                logger.info(f"OCR extracted {len(ocr_text)} characters from PDF")
            except Exception as e:
                logger.warning(f"OCR failed for PDF: {e}")
        elif file_ext in ['png', 'jpg', 'jpeg']:
            logger.info(f"Running OCR on image: {filename}")
            try:
                ocr_text = await ocr_service.extract_text_from_image(file_bytes)
                logger.info(f"OCR extracted {len(ocr_text)} characters from image")
            except Exception as e:
                logger.warning(f"OCR failed for image: {e}")
        elif file_ext in ['docx']:
            logger.info(f"Extracting text from DOCX: {filename}")
            try:
                import docx
                from io import BytesIO
                doc = docx.Document(BytesIO(file_bytes))
                
                # Extract paragraphs
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                
                # Extract text from tables
                table_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_text.append(" | ".join(row_text))
                
                # Combine all text
                all_text = paragraphs + table_text
                ocr_text = "\n".join(all_text)
                
                logger.info(f"Extracted {len(ocr_text)} characters from DOCX (including tables)")
            except Exception as e:
                logger.warning(f"Text extraction failed for DOCX: {e}")
        
        if not ocr_text:
            raise Exception(f"Failed to extract text from file: {filename}")
        
        # Update file record with OCR text
        await db["uploads"].update_one(
            {"file_id": file_id},
            {
                "$set": {
                    "ocr_text": ocr_text,
                    "processed": True,
                    "metadata.ocr_completed_at": datetime.utcnow()
                }
            }
        )
        
        # 3. Extract structured data using LLM
        structured_data = None
        try:
            logger.info(f"Extracting structured data from resume using LLM")
            structured_data = await llm_service.extract_resume_data(ocr_text)
            logger.info(f"Successfully extracted structured data: {list(structured_data.keys())}")
        except Exception as e:
            logger.warning(f"LLM data extraction failed: {e}")
        
        # 4. Update user profile with extracted data
        if structured_data:
            profile_updates = {}
            
            # Update skills if present
            if structured_data.get("skills"):
                profile_updates["profile.skills"] = structured_data["skills"]
            
            # Update experience if present
            if structured_data.get("experience"):
                profile_updates["profile.experience"] = structured_data["experience"]
            
            # Update education if present
            if structured_data.get("education"):
                profile_updates["profile.education"] = structured_data["education"]
            
            # Update certifications if present
            if structured_data.get("certifications"):
                profile_updates["profile.certifications"] = structured_data["certifications"]
            
            # Update contact info if present
            if structured_data.get("contact"):
                contact = structured_data["contact"]
                if contact.get("phone"):
                    profile_updates["profile.phone"] = contact["phone"]
                if contact.get("location"):
                    profile_updates["profile.location"] = contact["location"]
                if contact.get("linkedin"):
                    profile_updates["profile.linkedin_url"] = contact["linkedin"]
                if contact.get("github"):
                    profile_updates["profile.github_url"] = contact["github"]
                if contact.get("portfolio"):
                    profile_updates["profile.portfolio_url"] = contact["portfolio"]
            
            # Update summary if present
            if structured_data.get("summary"):
                profile_updates["profile.summary"] = structured_data["summary"]
            
            if profile_updates:
                profile_updates["updated_at"] = datetime.utcnow()
                
                await db["users"].update_one(
                    {"_id": user_id},
                    {"$set": profile_updates}
                )
                logger.info(f"Updated user profile with {len(profile_updates)} fields")
        
        # 5. Ingest into RAG system
        try:
            logger.info(f"Ingesting resume into RAG system")
            rag_service = RAGService(db, embeddings_service)
            
            doc_ids = await rag_service.ingest_document(
                user_id=user_id,
                content=ocr_text,
                doc_type="resume",
                metadata={
                    "source": "uploaded_file",
                    "file_id": file_id,
                    "filename": filename,
                    "uploaded_at": file_data["uploaded_at"].isoformat()
                }
            )
            
            logger.info(f"Ingested resume into RAG with {len(doc_ids)} chunks")
            
            # Update file record with RAG info
            await db["uploads"].update_one(
                {"file_id": file_id},
                {
                    "$set": {
                        "metadata.rag_doc_ids": doc_ids,
                        "metadata.rag_ingested_at": datetime.utcnow()
                    }
                }
            )
        except Exception as e:
            logger.warning(f"RAG ingestion failed: {e}")
        
        # Mark processing complete
        await db["uploads"].update_one(
            {"file_id": file_id},
            {
                "$set": {
                    "metadata.processing": False,
                    "metadata.processing_completed_at": datetime.utcnow()
                }
            }
        )
        
        logger.info(f"Resume processing completed successfully for file_id={file_id}")
        
        return {
            "status": "success",
            "file_id": file_id,
            "user_id": user_id,
            "ocr_text_length": len(ocr_text) if ocr_text else 0,
            "structured_data_extracted": structured_data is not None,
            "profile_updated": bool(profile_updates) if structured_data else False,
            "rag_chunks": len(doc_ids) if 'doc_ids' in locals() else 0,
            "processed_at": datetime.utcnow().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Resume processing failed for file_id={file_id}: {e}", exc_info=True)
        
        # Update file record with error
        try:
            await db["uploads"].update_one(
                {"file_id": file_id},
                {
                    "$set": {
                        "metadata.processing": False,
                        "metadata.processing_failed": True,
                        "metadata.error": str(e),
                        "metadata.failed_at": datetime.utcnow()
                    }
                }
            )
        except Exception as update_error:
            logger.error(f"Failed to update file status: {update_error}")
        
        # Retry logic
        if self.request.retries < self.max_retries:
            logger.info(f"Retrying task (attempt {self.request.retries + 1}/{self.max_retries})")
            raise self.retry(exc=e, countdown=120 * (self.request.retries + 1))
        
        return {
            "status": "failed",
            "file_id": file_id,
            "user_id": user_id,
            "error": str(e),
            "failed_at": datetime.utcnow().isoformat()
        }


@celery_app.task(name="cleanup_expired_resumes", bind=True)
@async_task
async def cleanup_expired_resumes(self, retention_days: int = 30) -> Dict[str, Any]:
    """
    Periodic task to clean up expired resume files from S3 and old uploads.
    
    Args:
        retention_days: Number of days to retain resumes (default: 30)
    
    Returns:
        Task result dictionary
    """
    db = get_task_db()
    
    try:
        logger.info(f"Starting cleanup of expired resumes (retention: {retention_days} days)")
        
        # Get worker services
        from app.workers.worker_services import get_worker_storage_service
        storage_service = get_worker_storage_service()
        
        # Calculate cutoff date
        cutoff_date = datetime.utcnow() - timedelta(days=retention_days)
        logger.info(f"Cleaning up resumes older than {cutoff_date}")
        
        deleted_resumes_count = 0
        deleted_uploads_count = 0
        deleted_s3_files = 0
        failed_deletions = []
        
        # 1. Find and clean up old resumes
        old_resumes = await db["resumes"].find({
            "generated_at": {"$lt": cutoff_date},
            "status": {"$in": [ResumeStatus.COMPLETED.value, ResumeStatus.FAILED.value]}
        }).to_list(length=1000)
        
        logger.info(f"Found {len(old_resumes)} expired resumes to clean up")
        
        for resume in old_resumes:
            resume_id = resume["resume_id"]
            s3_key = resume.get("s3_key")
            
            try:
                # Delete from S3 if exists
                if s3_key:
                    try:
                        await storage_service.delete_file(s3_key)
                        deleted_s3_files += 1
                        logger.debug(f"Deleted S3 file: {s3_key}")
                    except Exception as e:
                        logger.warning(f"Failed to delete S3 file {s3_key}: {e}")
                        failed_deletions.append({"type": "s3", "key": s3_key, "error": str(e)})
                
                # Delete from database or mark as archived
                await db["resumes"].update_one(
                    {"resume_id": resume_id},
                    {
                        "$set": {
                            "archived": True,
                            "archived_at": datetime.utcnow(),
                            "download_url": None,  # Remove expired presigned URL
                            "s3_key": None  # File deleted from S3
                        }
                    }
                )
                deleted_resumes_count += 1
                
            except Exception as e:
                logger.error(f"Failed to clean up resume {resume_id}: {e}")
                failed_deletions.append({"type": "resume", "resume_id": resume_id, "error": str(e)})
        
        # 2. Clean up old file uploads
        old_uploads = await db["uploads"].find({
            "uploaded_at": {"$lt": cutoff_date}
        }).to_list(length=1000)
        
        logger.info(f"Found {len(old_uploads)} expired uploads to clean up")
        
        for upload in old_uploads:
            file_id = upload["file_id"]
            s3_key = upload.get("s3_key")
            
            try:
                # Delete from S3
                if s3_key:
                    try:
                        await storage_service.delete_file(s3_key)
                        deleted_s3_files += 1
                        logger.debug(f"Deleted S3 upload: {s3_key}")
                    except Exception as e:
                        logger.warning(f"Failed to delete S3 upload {s3_key}: {e}")
                        failed_deletions.append({"type": "s3_upload", "key": s3_key, "error": str(e)})
                
                # Delete from database
                await db["uploads"].delete_one({"file_id": file_id})
                deleted_uploads_count += 1
                
            except Exception as e:
                logger.error(f"Failed to clean up upload {file_id}: {e}")
                failed_deletions.append({"type": "upload", "file_id": file_id, "error": str(e)})
        
        # 3. Clean up old failed/pending resumes (older than 7 days)
        failed_cutoff = datetime.utcnow() - timedelta(days=7)
        result = await db["resumes"].delete_many({
            "generated_at": {"$lt": failed_cutoff},
            "status": {"$in": [ResumeStatus.PENDING.value, ResumeStatus.PROCESSING.value]}
        })
        
        stale_count = result.deleted_count
        logger.info(f"Deleted {stale_count} stale pending/processing resumes")
        
        # 4. Clean up old audit logs (if configured)
        if hasattr(settings, 'AUDIT_LOG_RETENTION_DAYS'):
            audit_cutoff = datetime.utcnow() - timedelta(days=settings.AUDIT_LOG_RETENTION_DAYS)
            audit_result = await db["audit_logs"].delete_many({
                "timestamp": {"$lt": audit_cutoff}
            })
            audit_deleted = audit_result.deleted_count
            logger.info(f"Deleted {audit_deleted} old audit logs")
        else:
            audit_deleted = 0
        
        logger.info(
            f"Cleanup completed: {deleted_resumes_count} resumes archived, "
            f"{deleted_uploads_count} uploads deleted, {deleted_s3_files} S3 files deleted, "
            f"{stale_count} stale resumes deleted, {audit_deleted} audit logs deleted"
        )
        
        return {
            "status": "success",
            "deleted_resumes_count": deleted_resumes_count,
            "deleted_uploads_count": deleted_uploads_count,
            "deleted_s3_files": deleted_s3_files,
            "stale_resumes_deleted": stale_count,
            "audit_logs_deleted": audit_deleted,
            "failed_deletions": len(failed_deletions),
            "failed_deletions_details": failed_deletions[:10],  # First 10 failures
            "retention_days": retention_days,
            "cutoff_date": cutoff_date.isoformat(),
            "completed_at": datetime.utcnow().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Cleanup task failed: {e}", exc_info=True)
        
        return {
            "status": "failed",
            "error": str(e),
            "failed_at": datetime.utcnow().isoformat()
        }


# Configure periodic tasks
celery_app.conf.beat_schedule = {
    'cleanup-expired-resumes': {
        'task': 'cleanup_expired_resumes',
        'schedule': 86400.0,  # Once per day
    },
}
