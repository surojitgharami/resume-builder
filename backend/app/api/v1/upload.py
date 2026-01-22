# app/api/v1/upload.py
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Request
from motor.motor_asyncio import AsyncIOMotorDatabase
import uuid
from datetime import datetime
from io import BytesIO

from app.models.upload import FileUpload, UploadResponse, FileType
from app.models.user import User
from app.middleware.auth import get_current_active_user
from app.middleware.rate_limit import check_rate_limit
from app.db.mongo import get_database
from app.services.storage import get_storage_service
from app.services.ocr import get_ocr_service, OCRServiceUnavailable
from app.services.llm import get_llm_service
from app.services.embeddings import get_embeddings_service
from app.services.rag import RAGService
from app.core.config import settings
import logging

router = APIRouter()
logger = logging.getLogger(__name__)


@router.post("/upload", response_model=UploadResponse, status_code=status.HTTP_202_ACCEPTED)
async def upload_file(
    request: Request,
    file: UploadFile = File(...),
    file_type: FileType = FileType.RESUME,
    process_async: bool = True,
    current_user: User = Depends(get_current_active_user),
    db: AsyncIOMotorDatabase = Depends(get_database)
):
    """
    Upload a file (resume, cover letter, certificate, etc.).
    
    Can process file immediately or queue for background processing.
    
    Args:
        request: FastAPI request object
        file: Uploaded file
        file_type: Type of file being uploaded
        process_async: Whether to process file in background (default: True)
        current_user: Authenticated user
        db: Database connection
        
    Returns:
        Upload response with file details
        
    Raises:
        HTTPException: If upload fails or file is invalid
    """
    # Rate limiting
    await check_rate_limit(
        request,
        "file_upload",
        max_requests=20,
        window_seconds=3600
    )
    
    # Validate file extension
    file_ext = file.filename.split('.')[-1].lower()
    allowed_extensions = [ext.strip('.') for ext in settings.ALLOWED_UPLOAD_EXTENSIONS]
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type .{file_ext} not allowed. Allowed types: {', '.join(settings.ALLOWED_UPLOAD_EXTENSIONS)}"
        )
    
    # Check file size
    file_data = await file.read()
    file_size = len(file_data)
    max_size = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
    
    if file_size > max_size:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File size exceeds {settings.MAX_UPLOAD_SIZE_MB}MB limit"
        )
    
    try:
        storage_service = get_storage_service()
        
        # Generate unique file ID and S3 key
        file_id = str(uuid.uuid4())
        s3_key = f"uploads/{str(current_user.id)}/{file_type.value}/{file_id}.{file_ext}"
        
        # Upload to S3
        await storage_service.upload_file(
            BytesIO(file_data),
            s3_key,
            content_type=file.content_type,
            metadata={
                "user_id": str(current_user.id),
                "file_id": file_id,
                "file_type": file_type.value,
                "original_filename": file.filename
            }
        )
        
        # Process file (OCR if needed)
        ocr_text = None
        if file_type == FileType.RESUME:
            if process_async:
                # Queue background task for processing
                from app.workers.tasks import process_uploaded_resume
                
                task = process_uploaded_resume.delay(
                    file_id=file_id,
                    user_id=str(current_user.id)
                )
                
                # Update file record with task ID
                await uploads_coll.update_one(
                    {"file_id": file_id},
                    {
                        "$set": {
                            "metadata.processing_task_id": task.id,
                            "metadata.processing_queued_at": datetime.utcnow()
                        }
                    }
                )
                
                logger.info(f"Queued async file processing: file_id={file_id}, task_id={task.id}")
            else:
                # Process synchronously
                ocr_text = await process_resume_file(
                    file_data=file_data,
                    file_ext=file_ext,
                    user_id=str(current_user.id),
                    db=db
                )
        
        # Create file upload record
        file_upload = FileUpload(
            file_id=file_id,
            user_id=str(current_user.id),
            filename=file.filename,
            file_type=file_type,
            file_size=file_size,
            mime_type=file.content_type or "application/octet-stream",
            s3_key=s3_key,
            uploaded_at=datetime.utcnow(),
            processed=ocr_text is not None if not process_async else False,
            ocr_text=ocr_text,
            metadata={
                "original_filename": file.filename,
                "file_extension": file_ext,
                "process_async": process_async
            }
        )
        
        # Save to database
        uploads_coll = db["uploads"]
        await uploads_coll.insert_one(file_upload.model_dump())
        
        # Generate presigned download URL
        download_url = await storage_service.generate_presigned_url(
            s3_key,
            expiration=3600
        )
        
        logger.info(f"File uploaded successfully: {file_id}")
        
        return UploadResponse(
            file_id=file_id,
            filename=file.filename,
            file_size=file_size,
            uploaded_at=file_upload.uploaded_at,
            download_url=download_url,
            ocr_text=ocr_text
        )
        
    except Exception as e:
        logger.error(f"File upload failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"File upload failed: {str(e)}"
        )


async def process_resume_file(
    file_data: bytes,
    file_ext: str,
    user_id: str,
    db: AsyncIOMotorDatabase
) -> str:
    """
    Process uploaded resume file with OCR and data extraction.
    
    Args:
        file_data: File content as bytes
        file_ext: File extension
        user_id: User ID
        db: Database connection
        
    Returns:
        Extracted text from resume
    """
    try:
        ocr_service = get_ocr_service()
        
        # Check if OCR is available before attempting extraction
        if not ocr_service.is_available():
            status = ocr_service.get_availability_status()
            logger.warning(f"OCR service unavailable: {status['message']}")
            
            # For image/PDF files, we can't extract text without OCR
            if file_ext in ['pdf', 'png', 'jpg', 'jpeg']:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail={
                        "error": "OCR service unavailable",
                        "message": status['message'],
                        "suggestion": status.get('suggestion', 'Configure an OCR provider to process image and PDF files'),
                        "file_type": file_ext,
                        "workaround": "Please upload a DOCX file instead, or configure an OCR provider"
                    }
                )
        
        # Extract text based on file type
        if file_ext in ['pdf']:
            try:
                text = await ocr_service.extract_text_from_pdf(file_data)
            except OCRServiceUnavailable as e:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=e.to_dict()
                )
        elif file_ext in ['png', 'jpg', 'jpeg']:
            try:
                text = await ocr_service.extract_text_from_image(file_data)
            except OCRServiceUnavailable as e:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=e.to_dict()
                )
        elif file_ext in ['docx']:
            # Extract text from DOCX using python-docx
            try:
                import docx
                doc = docx.Document(BytesIO(file_data))
                
                # Extract all paragraphs
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                
                # Extract text from tables
                table_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_text.append(" | ".join(row_text))
                
                # Combine all text
                all_text = paragraphs + table_text
                text = "\n".join(all_text)
                
                logger.info(f"Extracted {len(text)} characters from DOCX file")
            except Exception as e:
                logger.error(f"DOCX extraction failed: {e}")
                raise Exception(f"Failed to extract text from DOCX: {str(e)}")
        else:
            logger.warning(f"Text extraction not supported for .{file_ext}")
            return None
        
        # Extract structured data using LLM
        llm_service = get_llm_service()
        try:
            structured_data = await llm_service.extract_resume_data(text)
            
            # Update user profile with extracted data
            users_coll = db["users"]
            await users_coll.update_one(
                {"_id": user_id},
                {
                    "$set": {
                        "profile.skills": structured_data.get("skills", []),
                        "profile.experience": structured_data.get("experience", []),
                        "profile.education": structured_data.get("education", []),
                        "profile.certifications": structured_data.get("certifications", []),
                        "updated_at": datetime.utcnow()
                    }
                }
            )
            logger.info(f"Updated user profile from resume extraction")
        except Exception as e:
            logger.warning(f"Failed to extract structured data: {e}")
        
        # Ingest into RAG system
        try:
            embeddings_service = get_embeddings_service()
            rag_service = RAGService(db, embeddings_service)
            
            await rag_service.ingest_document(
                user_id=user_id,
                content=text,
                doc_type="resume",
                metadata={"source": "uploaded_file"}
            )
            logger.info(f"Resume ingested into RAG system")
        except Exception as e:
            logger.warning(f"Failed to ingest into RAG: {e}")
        
        return text
        
    except Exception as e:
        logger.error(f"Resume processing failed: {e}")
        return None


@router.get("/uploads", response_model=list[UploadResponse])
async def list_uploads(
    file_type: FileType = None,
    limit: int = 10,
    skip: int = 0,
    current_user: User = Depends(get_current_active_user),
    db: AsyncIOMotorDatabase = Depends(get_database)
):
    """
    Get all uploaded files for the current user.
    
    Args:
        file_type: Optional filter by file type
        limit: Maximum number of files to return
        skip: Number of files to skip
        current_user: Authenticated user
        db: Database connection
        
    Returns:
        List of upload responses
    """
    uploads_coll = db["uploads"]
    
    query = {"user_id": str(current_user.id)}
    if file_type:
        query["file_type"] = file_type.value
    
    cursor = uploads_coll.find(query).sort("uploaded_at", -1).skip(skip).limit(limit)
    uploads = await cursor.to_list(length=limit)
    
    storage_service = get_storage_service()
    
    responses = []
    for upload in uploads:
        # Generate fresh download URL
        try:
            download_url = await storage_service.generate_presigned_url(
                upload["s3_key"],
                expiration=3600
            )
        except:
            download_url = None
        
        responses.append(UploadResponse(
            file_id=upload["file_id"],
            filename=upload["filename"],
            file_size=upload["file_size"],
            uploaded_at=upload["uploaded_at"],
            download_url=download_url,
            ocr_text=upload.get("ocr_text")
        ))
    
    return responses


@router.delete("/uploads/{file_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_upload(
    file_id: str,
    current_user: User = Depends(get_current_active_user),
    db: AsyncIOMotorDatabase = Depends(get_database)
):
    """
    Delete an uploaded file.
    
    Args:
        file_id: File ID
        current_user: Authenticated user
        db: Database connection
        
    Raises:
        HTTPException: If file not found
    """
    uploads_coll = db["uploads"]
    
    upload = await uploads_coll.find_one({
        "file_id": file_id,
        "user_id": str(current_user.id)
    })
    
    if not upload:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="File not found"
        )
    
    # Delete from S3
    try:
        storage_service = get_storage_service()
        await storage_service.delete_file(upload["s3_key"])
    except Exception as e:
        logger.warning(f"Failed to delete S3 file: {e}")
    
    # Delete from database
    await uploads_coll.delete_one({"file_id": file_id})
    
    return None
